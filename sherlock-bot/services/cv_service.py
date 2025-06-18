# services/cv_service.py - COMPLETELY FIXED CV analysis service
import os
import uuid
import time
import requests
import PyPDF2
import docx
import traceback
from datetime import datetime
import re
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem, Table, TableStyle
from firebase_admin import storage
from services.firebase_service import download_file_from_storage, get_file_download_url
from utils.logger import get_logger
from config import Config

# Initialize logger
logger = get_logger()

# FIXED: Handle NLTK properly with fallbacks
try:
    import nltk
    from nltk.tokenize import sent_tokenize
    
    # Try to download required NLTK data with better error handling
    try:
        # Check if punkt_tab exists (newer NLTK versions)
        nltk.data.find('tokenizers/punkt_tab')
        logger.info("✅ Found punkt_tab tokenizer")
    except LookupError:
        try:
            # Try to download punkt_tab
            nltk.download('punkt_tab', quiet=True)
            logger.info("✅ Downloaded punkt_tab tokenizer")
        except Exception:
            try:
                # Fallback to older punkt
                nltk.data.find('tokenizers/punkt')
                logger.info("✅ Found punkt tokenizer")
            except LookupError:
                try:
                    nltk.download('punkt', quiet=True)
                    logger.info("✅ Downloaded punkt tokenizer")
                except Exception as e:
                    logger.warning(f"⚠️ Could not download NLTK data: {e}")
                    # Define fallback function
                    def sent_tokenize(text):
                        """Fallback sentence tokenizer"""
                        sentences = re.split(r'[.!?]+', text)
                        return [s.strip() for s in sentences if s.strip()]
                        
except ImportError:
    logger.warning("⚠️ NLTK not available, using regex fallback")
    def sent_tokenize(text):
        """Basic sentence tokenizer fallback"""
        sentences = re.split(r'[.!?]+', text)
        return [s.strip() for s in sentences if s.strip()]


def process_basic_review(storage_path):
    """Process basic CV review with comprehensive error handling"""
    try:
        logger.info(f"🔄 Starting basic review for: {storage_path}")
        
        # Download file from storage
        local_file_path = download_file_from_storage(storage_path)
        logger.info(f"📥 Downloaded file to: {local_file_path}")
        
        # Extract CV data
        cv_data = extract_text_from_cv(local_file_path)
        cv_data['file_path'] = local_file_path

        # Process review
        if Config.CV_ANALYSIS_API_URL:
            logger.info("🔗 Using external CV analysis API")
            review_result = call_cv_analysis_api(cv_data, 'basic')
        else:
            logger.info("🤖 Using internal CV analysis")
            review_result = analyze_cv_basic(cv_data)

        # Add metadata
        review_result['cv_file_name'] = os.path.basename(storage_path)
        review_result['review_type'] = 'basic'
        review_result['success'] = True

        # Clean up temporary file
        if os.path.exists(local_file_path):
            os.remove(local_file_path)
            logger.info(f"🗑️ Cleaned up: {local_file_path}")

        logger.info("✅ Basic review completed successfully")
        return review_result

    except Exception as e:
        logger.error(f"❌ Error in basic review processing: {str(e)}")
        return {'success': False, 'error': str(e)}


def process_advanced_review(storage_path):
    """Process advanced CV review"""
    try:
        logger.info(f"🔄 Starting advanced review for: {storage_path}")
        
        local_file_path = download_file_from_storage(storage_path)
        cv_data = extract_text_from_cv(local_file_path)
        cv_data['file_path'] = local_file_path

        if Config.CV_ANALYSIS_API_URL:
            review_result = call_cv_analysis_api(cv_data, 'advanced')
        else:
            review_result = analyze_cv_advanced(cv_data)

        # Generate PDF report
        report_path = generate_pdf_report(review_result, local_file_path)

        # Upload report to Firebase Storage
        path_parts = storage_path.split('/')
        phone_number = path_parts[1] if len(path_parts) > 1 else 'unknown'

        report_folder = 'review-reports'
        report_filename = f"report_{int(time.time())}_{uuid.uuid4().hex[:8]}.pdf"
        report_storage_path = f"{report_folder}/{phone_number}/{report_filename}"

        bucket = storage.bucket()
        blob = bucket.blob(report_storage_path)
        blob.upload_from_filename(report_path)
        blob.content_type = 'application/pdf'

        download_url = get_file_download_url(report_storage_path)
        review_result['download_link'] = download_url

        review_result['cv_file_name'] = os.path.basename(storage_path)
        review_result['review_type'] = 'advanced'
        review_result['report_path'] = report_storage_path
        review_result['success'] = True

        # Clean up temporary files
        if os.path.exists(local_file_path):
            os.remove(local_file_path)
        if os.path.exists(report_path):
            os.remove(report_path)

        logger.info("✅ Advanced review completed successfully")
        return review_result

    except Exception as e:
        logger.error(f"❌ Error in advanced review processing: {str(e)}")
        return {'success': False, 'error': str(e)}


def extract_text_from_cv(file_path):
    """Extract text from CV with better error handling"""
    try:
        ext = os.path.splitext(file_path)[1].lower()
        logger.info(f"📄 Processing file type: {ext}")

        if ext == '.pdf':
            return extract_text_from_pdf(file_path)
        elif ext in ['.docx', '.doc']:
            return extract_text_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")

    except Exception as e:
        logger.error(f"❌ Error extracting text from {file_path}: {str(e)}")
        # Return minimal fallback structure
        return {
            'full_text': f"Error extracting text: {str(e)}",
            'sections': {},
            'contact_info': {},
            'metrics': {'word_count': 0, 'sentence_count': 0, 'avg_sentence_length': 0, 'bullet_points': 0},
            'sentences': []
        }


def extract_text_from_pdf(file_path):
    """Extract text from PDF"""
    try:
        text = ""
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            num_pages = len(reader.pages)

            for page_num in range(num_pages):
                page = reader.pages[page_num]
                text += page.extract_text() + "\n"

        cv_data = analyze_cv_structure(text)
        cv_data['full_text'] = text
        cv_data['metadata'] = {
            'page_count': num_pages,
            'file_type': 'pdf'
        }

        logger.info(f"📊 Extracted {len(text)} characters from PDF")
        return cv_data

    except Exception as e:
        logger.error(f"❌ Error extracting text from PDF {file_path}: {str(e)}")
        raise e


def extract_text_from_docx(file_path):
    """Extract text from DOCX with better error handling"""
    try:
        logger.info(f"📝 Extracting text from DOCX: {file_path}")
        doc = docx.Document(file_path)
        text = ""

        for para in doc.paragraphs:
            text += para.text + "\n"

        logger.info(f"📊 Extracted {len(text)} characters from DOCX")
        
        cv_data = analyze_cv_structure(text)
        cv_data['full_text'] = text
        cv_data['metadata'] = {
            'paragraph_count': len(doc.paragraphs),
            'file_type': 'docx'
        }

        return cv_data

    except Exception as e:
        logger.error(f"❌ Error extracting text from DOCX {file_path}: {str(e)}")
        raise e


def analyze_cv_structure(text):
    """Analyze CV structure with NLTK fallback"""
    try:
        # Use sentence tokenizer (NLTK or fallback)
        sentences = sent_tokenize(text)
        logger.info(f"📝 Found {len(sentences)} sentences")
        
    except Exception as e:
        logger.warning(f"⚠️ Sentence tokenization failed, using regex fallback: {str(e)}")
        sentences = re.split(r'[.!?]+', text)
        sentences = [s.strip() for s in sentences if s.strip()]
    
    sections = identify_sections(text)
    contact_info = {
        'email': extract_email(text),
        'phone': extract_phone(text),
        'linkedin': extract_linkedin(text)
    }
    
    word_count = len(text.split())
    sentence_count = len(sentences)
    avg_sentence_length = word_count / max(sentence_count, 1)
    
    metrics = {
        'word_count': word_count,
        'sentence_count': sentence_count,
        'avg_sentence_length': avg_sentence_length,
        'bullet_points': len(re.findall(r'•|\*|-', text))
    }

    logger.info(f"📊 Analysis complete: {word_count} words, {sentence_count} sentences")

    return {
        'sections': sections,
        'contact_info': contact_info,
        'metrics': metrics,
        'sentences': sentences
    }


def identify_sections(text):
    """Identify CV sections"""
    sections = {}
    section_patterns = {
        'summary': r'(profile|summary|objective|about\s*me)',
        'experience': r'(experience|employment|work\s*history|professional\s*background)',
        'education': r'(education|qualification|academic|degree|university)',
        'skills': r'(skills|expertise|competencies|proficiencies|technical)',
        'projects': r'(projects|portfolio|works)',
        'certifications': r'(certifications|certificates|credentials)',
        'languages': r'(languages|language\s*proficiency)',
        'interests': r'(interests|hobbies|activities)'
    }

    lines = text.split('\n')
    current_section = None
    section_content = {}

    for line in lines:
        line = line.strip()
        if not line:
            continue

        found_section = None
        for section_name, pattern in section_patterns.items():
            # Fixed: Remove (?i) from pattern since we're using re.IGNORECASE flag
            if re.search(f"^{pattern}.*$", line, re.IGNORECASE):
                found_section = section_name
                break

        if found_section:
            current_section = found_section
            section_content[current_section] = []
        elif current_section:
            section_content[current_section].append(line)

    for section_name, lines in section_content.items():
        sections[section_name] = '\n'.join(lines)

    return sections


def extract_email(text):
    """Extract email address"""
    email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
    match = re.search(email_pattern, text)
    return match.group(0) if match else ''


def extract_phone(text):
    """Extract phone number"""
    phone_pattern = r'(?:\+\d{1,3}[\s-]?)?\(?\d{3,4}\)?[\s.-]?\d{3}[\s.-]?\d{4}'
    match = re.search(phone_pattern, text)
    return match.group(0) if match else ''


def extract_linkedin(text):
    """Extract LinkedIn URL"""
    linkedin_pattern = r'(?:https?://)?(?:www\.)?linkedin\.com/in/[a-zA-Z0-9_-]+'
    match = re.search(linkedin_pattern, text)
    return match.group(0) if match else ''


def analyze_cv_basic(cv_data):
    """Basic CV analysis with fallback insights"""
    try:
        text = cv_data.get('full_text', '')
        sections = cv_data.get('sections', {})
        metrics = cv_data.get('metrics', {})
        contact_info = cv_data.get('contact_info', {})

        insights = []

        # Check for essential sections
        if 'summary' not in sections and 'profile' not in text.lower():
            insights.append("Consider adding a professional summary at the top of your CV to highlight your key qualifications.")

        if 'experience' not in sections and 'work' not in text.lower():
            insights.append("Your work experience section is missing or not clearly defined. This is a critical section for most CVs.")

        if 'education' not in sections and 'university' not in text.lower():
            insights.append("Include your educational background with relevant details about degrees, institutions, and graduation dates.")

        if 'skills' not in sections and 'skill' not in text.lower():
            insights.append("A skills section would help highlight your key competencies relevant to your target roles.")

        # Check contact information
        if not contact_info.get('email') and '@' not in text:
            insights.append("Ensure your contact information including email is clearly visible at the top of your CV.")

        # Check for quantifiable achievements
        if not any(char in text for char in ['%', '₦', '$', '+', 'increase', 'improve', 'reduce']):
            insights.append("Add quantifiable achievements with metrics (%, numbers, etc.) to make your accomplishments more impactful.")

        # Check formatting
        if metrics.get('bullet_points', 0) < 5:
            insights.append("Use bullet points to make your CV more scannable and highlight key accomplishments.")

        # Check word count
        word_count = metrics.get('word_count', 0)
        if word_count < 200:
            insights.append("Your CV appears to be quite short. Consider adding more details about your experience and skills.")
        elif word_count > 1000:
            insights.append("Your CV may be too lengthy. Consider condensing it to 1-2 pages for better readability.")

        # Add general tips if we don't have enough specific insights
        general_insights = [
            "Use action verbs at the beginning of bullet points to create a stronger impression.",
            "Ensure consistent formatting throughout your CV for a professional appearance.",
            "Tailor your CV for each job application to highlight the most relevant experience.",
            "Proofread carefully for grammar and spelling errors.",
            "Use industry-specific keywords to pass through automated screening systems."
        ]

        # Ensure we have at least 5 insights
        while len(insights) < 5 and general_insights:
            insights.append(general_insights.pop(0))

        logger.info(f"✅ Generated {len(insights)} insights for basic review")

        return {
            'success': True,
            'timestamp': datetime.now().isoformat(),
            'insights': insights[:8],  # Limit to 8 insights
            'api_provider': 'Internal Analysis'
        }

    except Exception as e:
        logger.error(f"❌ Error in basic CV analysis: {str(e)}")
        return {
            'success': False,
            'error': str(e)
        }


def analyze_cv_advanced(cv_data):
    """Advanced CV analysis - simplified version"""
    try:
        # For now, use basic analysis with additional scoring
        basic_result = analyze_cv_basic(cv_data)
        
        if not basic_result.get('success'):
            return basic_result
        
        # Add scoring and additional insights for advanced review
        text = cv_data.get('full_text', '')
        sections = cv_data.get('sections', {})
        metrics = cv_data.get('metrics', {})
        
        # Calculate improvement score
        score = 60  # Base score
        
        # Add points for sections present
        essential_sections = ['summary', 'experience', 'education', 'skills']
        for section in essential_sections:
            if section in sections:
                score += 5
        
        # Add points for good content
        word_count = metrics.get('word_count', 0)
        if 300 <= word_count <= 1000:
            score += 10
        
        # Add points for metrics/achievements
        if any(char in text for char in ['%', '₦', '$', '+', 'increase', 'improve']):
            score += 10
        
        # Cap the score
        score = min(max(score, 40), 95)
        
        # Enhanced insights for advanced review
        insights = basic_result['insights']
        insights.extend([
            "FORMATTING: Maintain consistent formatting with a clear hierarchy throughout your document.",
            "KEYWORDS: Include more industry-specific keywords to pass through applicant tracking systems (ATS).",
            "RELEVANCE: Focus on your most recent and relevant experience for your target roles."
        ])
        
        return {
            'success': True,
            'timestamp': datetime.now().isoformat(),
            'improvement_score': score,
            'insights': insights[:10],
            'api_provider': 'Internal Analysis'
        }
        
    except Exception as e:
        logger.error(f"❌ Error in advanced CV analysis: {str(e)}")
        return {
            'success': False,
            'error': str(e)
        }


def call_cv_analysis_api(cv_data, review_type):
    """Call external CV analysis API with correct parameters"""
    try:
        local_file_path = cv_data.get('file_path', None)
        if not local_file_path:
            logger.error("File path missing from CV data")
            return analyze_cv_fallback(cv_data, review_type)

        logger.info(f"📤 Calling CV API with file: {os.path.basename(local_file_path)}")

        # Open file for upload
        with open(local_file_path, 'rb') as f:
            files = {
                'cv': (os.path.basename(local_file_path), f, 'application/octet-stream')
            }

            # FIXED: Use correct form data with all three required fields
            form_data = {
                'job_title': 'General Application',  # Can be customized based on user input
                'job_description': 'Seeking opportunities in various industries. Review CV for general job applications including corporate, technical, and professional roles.'
            }

            api_url = Config.CV_ANALYSIS_API_URL
            logger.info(f"🌐 API URL: {api_url}")
            
            response = requests.post(
                api_url,
                files=files,
                data=form_data,
                timeout=60
            )

        logger.info(f"📥 API Response Status: {response.status_code}")
        
        if response.status_code != 200:
            logger.error(f"API returned status {response.status_code}: {response.text[:500]}")
            return analyze_cv_fallback(cv_data, review_type)

        try:
            result = response.json()
            logger.info(f"📊 API Response Keys: {list(result.keys()) if isinstance(result, dict) else 'Not a dict'}")
            logger.info(f"📝 First 200 chars of response: {str(result)[:200]}")
        except Exception as json_error:
            logger.error(f"Failed to parse API JSON response: {json_error}")
            logger.error(f"Raw response: {response.text[:500]}")
            return analyze_cv_fallback(cv_data, review_type)

        if review_type == 'basic':
            processed_result = process_basic_api_response(result)
        else:
            processed_result = process_advanced_api_response(result)
            
        logger.info(f"✅ Processed result has {len(processed_result.get('insights', []))} insights")
        
        # If no insights from API, use fallback
        if not processed_result.get('insights'):
            logger.warning("⚠️ No insights from API, using fallback")
            return analyze_cv_fallback(cv_data, review_type)
            
        return processed_result

    except Exception as e:
        logger.error(f"Error calling CV analysis API: {str(e)}")
        import traceback
        logger.error(f"Full error: {traceback.format_exc()}")
        return analyze_cv_fallback(cv_data, review_type)


def analyze_cv_fallback(cv_data, review_type):
    """Fallback analysis"""
    if review_type == 'basic':
        return analyze_cv_basic(cv_data)
    else:
        return analyze_cv_advanced(cv_data)


def process_basic_api_response(result):
    """Process basic API response with better error handling"""
    insights = []
    
    # Check if API returned an error
    if result.get('error'):
        logger.error(f"API Error: {result.get('error')}")
        return {
            'success': False,
            'error': result.get('error')
        }
    
    # Log the response structure for debugging
    logger.info(f"API Response structure: {list(result.keys()) if isinstance(result, dict) else type(result)}")
    
    # Extract insights from section feedback
    section_feedback = result.get('section_feedback', {})
    if section_feedback:
        for section, details in section_feedback.items():
            if isinstance(details, dict):
                if details.get('improvement_needed', False):
                    suggestion = details.get('suggestion', '') or details.get('feedback', '')
                    if suggestion:
                        insights.append(f"{section.upper()}: {suggestion}")
            elif isinstance(details, str) and details:
                # Handle case where details is just a string
                insights.append(f"{section.upper()}: {details}")
    
    # Extract general suggestions
    general_suggestions = result.get('general_suggestions', [])
    if isinstance(general_suggestions, list):
        for suggestion in general_suggestions:
            if suggestion:
                insights.append(suggestion)
    
    # Try alternative field names that the API might use
    alternative_fields = ['feedback', 'suggestions', 'improvements', 'recommendations']
    for field in alternative_fields:
        if field in result and isinstance(result[field], list):
            for item in result[field]:
                if item and item not in insights:
                    insights.append(item)
    
    # If still no insights, check for a general analysis or message field
    if not insights:
        if 'analysis' in result and result['analysis']:
            insights.append(result['analysis'])
        if 'message' in result and result['message']:
            insights.append(result['message'])
    
    logger.info(f"📊 Extracted {len(insights)} insights from API response")
    
    # If no insights found, return failure so fallback is used
    if not insights:
        logger.warning("No insights found in API response, will use fallback")
        return {
            'success': False,
            'error': 'No insights returned from API'
        }

    return {
        'success': True,
        'timestamp': datetime.now().isoformat(),
        'insights': insights[:8],
        'api_provider': 'CV Analyzer API'
    }


def process_advanced_api_response(result):
    """Process advanced API response with better error handling"""
    # Check if API returned an error
    if result.get('error'):
        logger.error(f"API Error: {result.get('error')}")
        return {
            'success': False,
            'error': result.get('error')
        }
    
    # Log the response structure for debugging
    logger.info(f"API Response structure: {list(result.keys()) if isinstance(result, dict) else type(result)}")
    
    # Extract score (try different possible field names)
    score = result.get('overall_score') or result.get('score') or result.get('cv_score') or 65
    section_scores = {}

    # Extract section scores
    section_feedback = result.get('section_feedback', {})
    if section_feedback:
        for section, details in section_feedback.items():
            if isinstance(details, dict):
                section_scores[section] = details.get('score', 50)

    # Extract insights
    insights = []
    
    # From section feedback
    if section_feedback:
        for section, details in section_feedback.items():
            if isinstance(details, dict):
                suggestion = details.get('suggestion', '') or details.get('feedback', '')
                if suggestion:
                    insights.append(f"{section.upper()}: {suggestion}")
            elif isinstance(details, str) and details:
                insights.append(f"{section.upper()}: {details}")
    
    # Add keyword analysis if available
    if 'keyword_matches' in result:
        keyword_percentage = result.get('keyword_match_percentage', 0)
        insights.append(f"KEYWORDS: Your CV matches {keyword_percentage}% of the job keywords. Consider adding more relevant keywords.")
    elif 'keywords' in result:
        insights.append(f"KEYWORDS: {result['keywords']}")
    
    # Try to get general suggestions
    general_suggestions = result.get('general_suggestions', [])
    if isinstance(general_suggestions, list):
        insights.extend(general_suggestions)
    
    # Try alternative field names
    alternative_fields = ['feedback', 'suggestions', 'improvements', 'recommendations']
    for field in alternative_fields:
        if field in result and isinstance(result[field], list):
            for item in result[field]:
                if item and item not in insights:
                    insights.append(item)
    
    logger.info(f"📊 Extracted {len(insights)} insights from API response")
    
    # If no insights found, return failure so fallback is used
    if not insights:
        logger.warning("No insights found in API response, will use fallback")
        return {
            'success': False,
            'error': 'No insights returned from API'
        }

    return {
        'success': True,
        'timestamp': datetime.now().isoformat(),
        'improvement_score': score,
        'section_scores': section_scores,
        'insights': insights[:10],  # Allow more insights for advanced review
        'api_provider': 'CV Analyzer API'
    }


def generate_pdf_report(review_result, cv_path):
    """Generate PDF report - simplified version"""
    try:
        timestamp = int(time.time())
        report_filename = f"cv_review_report_{timestamp}.pdf"
        
        # Create uploads directory if it doesn't exist
        uploads_dir = '/tmp/uploads' if os.path.exists('/tmp') else 'uploads'
        os.makedirs(uploads_dir, exist_ok=True)
        report_path = os.path.join(uploads_dir, report_filename)

        doc = SimpleDocTemplate(
            report_path,
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72
        )

        styles = getSampleStyleSheet()
        content = []

        # Title
        content.append(Paragraph("CV Review Report", styles['Title']))
        content.append(Spacer(1, 12))

        # Date
        date_text = f"Generated on: {datetime.now().strftime('%d %B %Y')}"
        content.append(Paragraph(date_text, styles['Normal']))
        content.append(Spacer(1, 24))

        # Score (if available)
        score = review_result.get('improvement_score')
        if score:
            content.append(Paragraph(f"CV Improvement Score: {score}/100", styles['Heading1']))
            content.append(Spacer(1, 12))

        # Insights
        content.append(Paragraph("Key Insights and Recommendations", styles['Heading1']))
        content.append(Spacer(1, 12))

        insights = review_result.get('insights', [])
        for i, insight in enumerate(insights, 1):
            content.append(Paragraph(f"{i}. {insight}", styles['Normal']))
            content.append(Spacer(1, 6))

        doc.build(content)
        logger.info(f"✅ Generated PDF report: {report_path}")
        return report_path

    except Exception as e:
        logger.error(f"❌ Error generating PDF report: {str(e)}")
        raise e