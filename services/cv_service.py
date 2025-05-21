# services/cv_service.py - CV analysis service
import os
import uuid
import time
import requests
import PyPDF2
import docx
from datetime import datetime
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem
from services.firebase_service import download_file_from_storage, get_file_download_url
from services.firestore_service import save_review_result, get_user_email
from services.sendgrid_service import send_review_email
from utils.logger import get_logger
from config import Config

# Initialize logger
logger = get_logger()

def process_basic_review(storage_path):
    """
    Process CV for basic review
    
    Args:
        storage_path (str): Path to CV file in Firebase Storage
        
    Returns:
        dict: Review results
    """
    try:
        # Download file from Firebase Storage to temp location
        local_file_path = download_file_from_storage(storage_path)
        
        # Extract text from CV
        cv_data = extract_text_from_cv(local_file_path)
        
        # Get phone number
        phone_number = get_phone_number_from_storage_path(storage_path)
        
        # Call CV analysis API or use fallback
        if Config.CV_ANALYSIS_API_URL and Config.CV_ANALYSIS_API_KEY:
            review_result = call_cv_analysis_api(cv_data, 'basic')
        else:
            review_result = generate_basic_review_fallback(cv_data)
        
        # Add metadata
        review_result['cv_file_name'] = os.path.basename(storage_path)
        review_result['review_type'] = 'basic'
        
        # Success flag
        review_result['success'] = True
        
        # Clean up the local temp file
        if os.path.exists(local_file_path):
            os.remove(local_file_path)
        
        return review_result
    
    except Exception as e:
        logger.error(f"Error in basic review processing: {str(e)}")
        return {
            'success': False,
            'error': str(e)
        }

def process_advanced_review(storage_path):
    """
    Process CV for advanced review
    
    Args:
        storage_path (str): Path to CV file in Firebase Storage
        
    Returns:
        dict: Review results
    """
    try:
        # Download file from Firebase Storage to temp location
        local_file_path = download_file_from_storage(storage_path)
        
        # Extract text from CV
        cv_data = extract_text_from_cv(local_file_path)
        
        # Get phone number
        phone_number = get_phone_number_from_storage_path(storage_path)
        
        # Call CV analysis API or use fallback
        if Config.CV_ANALYSIS_API_URL and Config.CV_ANALYSIS_API_KEY:
            review_result = call_cv_analysis_api(cv_data, 'advanced')
        else:
            review_result = generate_advanced_review_fallback(cv_data)
        
        # Generate a PDF report
        report_path = generate_pdf_report(review_result, local_file_path)
        
        # Upload report to Firebase Storage
        report_folder = 'review-reports'
        report_filename = f"report_{int(time.time())}_{uuid.uuid4().hex[:8]}.pdf"
        report_storage_path = f"{report_folder}/{phone_number}/{report_filename}"
        
        bucket = storage.bucket()
        blob = bucket.blob(report_storage_path)
        blob.upload_from_filename(report_path)
        blob.content_type = 'application/pdf'
        
        # Get a download URL for the report
        download_url = get_file_download_url(report_storage_path)
        review_result['download_link'] = download_url
        
        # Add metadata
        review_result['cv_file_name'] = os.path.basename(storage_path)
        review_result['review_type'] = 'advanced'
        review_result['report_path'] = report_storage_path
        
        # Success flag
        review_result['success'] = True
        
        # Clean up the local temp files
        if os.path.exists(local_file_path):
            os.remove(local_file_path)
        if os.path.exists(report_path):
            os.remove(report_path)
        
        return review_result
    
    except Exception as e:
        logger.error(f"Error in advanced review processing: {str(e)}")
        return {
            'success': False,
            'error': str(e)
        }

def extract_text_from_cv(file_path):
    """
    Extract text and structure from a CV file
    
    Args:
        file_path (str): Path to the CV file
        
    Returns:
        dict: Extracted CV data with text and structure
    """
    try:
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == '.pdf':
            return extract_text_from_pdf(file_path)
        elif ext in ['.docx', '.doc']:
            return extract_text_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")
    
    except Exception as e:
        logger.error(f"Error extracting text from {file_path}: {str(e)}")
        # Return basic text extraction as fallback
        with open(file_path, 'rb') as f:
            text = str(f.read())
        return {'full_text': text}

def extract_text_from_pdf(file_path):
    """
    Extract text from a PDF file with structure recognition
    
    Args:
        file_path (str): Path to the PDF file
        
    Returns:
        dict: Extracted data with text and structure
    """
    try:
        text = ""
        
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            num_pages = len(reader.pages)
            
            # Extract text from all pages
            for page_num in range(num_pages):
                page = reader.pages[page_num]
                text += page.extract_text() + "\n"
        
        # Process the text to identify sections
        sections = identify_sections(text)
        
        return {
            'full_text': text,
            'sections': sections,
            'metadata': {
                'page_count': num_pages,
                'file_type': 'pdf'
            }
        }
    
    except Exception as e:
        logger.error(f"Error extracting text from PDF {file_path}: {str(e)}")
        raise e

def extract_text_from_docx(file_path):
    """
    Extract text from a DOCX file with structure recognition
    
    Args:
        file_path (str): Path to the DOCX file
        
    Returns:
        dict: Extracted data with text and structure
    """
    try:
        doc = docx.Document(file_path)
        text = ""
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            text += para.text + "\n"
        
        # Process the text to identify sections
        sections = identify_sections(text)
        
        return {
            'full_text': text,
            'sections': sections,
            'metadata': {
                'paragraph_count': len(doc.paragraphs),
                'file_type': 'docx'
            }
        }
    
    except Exception as e:
        logger.error(f"Error extracting text from DOCX {file_path}: {str(e)}")
        raise e

def identify_sections(text):
    """
    Identify common CV sections in text
    
    Args:
        text (str): CV text
        
    Returns:
        dict: Identified sections
    """
    import re
    
    sections = {}
    
    # Define section patterns
    section_patterns = {
        'summary': r'(?i)(profile|summary|objective|about\s*me)',
        'experience': r'(?i)(experience|employment|work\s*history|professional\s*background)',
        'education': r'(?i)(education|qualification|academic|degree|university)',
        'skills': r'(?i)(skills|expertise|competencies|proficiencies|technical)',
        'projects': r'(?i)(projects|portfolio|works)',
        'certifications': r'(?i)(certifications|certificates|credentials)',
        'languages': r'(?i)(languages|language\s*proficiency)',
        'interests': r'(?i)(interests|hobbies|activities)'
    }
    
    # Find each section in the text
    for section_name, pattern in section_patterns.items():
        match = re.search(f"({pattern}).*?\n", text, re.IGNORECASE)
        if match:
            start_pos = match.end()
            
            # Find the next section heading
            next_section_pattern = '|'.join(f"({p})" for p in section_patterns.values())
            next_match = re.search(f"({next_section_pattern}).*?\n", text[start_pos:], re.IGNORECASE)
            
            if next_match:
                end_pos = start_pos + next_match.start()
                sections[section_name] = text[start_pos:end_pos].strip()
            else:
                sections[section_name] = text[start_pos:].strip()
    
    # Extract contact information
    contact_info = {
        'email': extract_email(text),
        'phone': extract_phone(text)
    }
    
    return {
        'sections': sections,
        'contact_info': contact_info
    }

def extract_email(text):
    """Extract email from text"""
    import re
    email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
    match = re.search(email_pattern, text)
    return match.group(0) if match else ''

def extract_phone(text):
    """Extract phone number from text"""
    import re
    phone_pattern = r'(?:\+\d{1,3}[\s-]?)?\(?\d{3,4}\)?[\s.-]?\d{3}[\s.-]?\d{4}'
    match = re.search(phone_pattern, text)
    return match.group(0) if match else ''

def call_cv_analysis_api(cv_data, review_type):
    """
    Call external CV analysis API
    
    Args:
        cv_data (dict): CV data with text and structure
        review_type (str): Type of review (basic or advanced)
        
    Returns:
        dict: API response with review results
    """
    try:
        # Prepare request data
        request_data = {
            'cv_text': cv_data.get('full_text', ''),
            'sections': cv_data.get('sections', {}),
            'review_type': review_type
        }
        
        # Call the API
        headers = {
            'Content-Type': 'application/json',
            'Authorization': f'Bearer {Config.CV_ANALYSIS_API_KEY}'
        }
        
        response = requests.post(
            Config.CV_ANALYSIS_API_URL,
            json=request_data,
            headers=headers,
            timeout=30  # 30 second timeout
        )
        
        # Check for successful response
        if response.status_code != 200:
            logger.error(f"API returned status {response.status_code}: {response.text}")
            return generate_review_fallback(cv_data, review_type)
        
        # Parse response
        result = response.json()
        
        # Validate API response
        if not result.get('success'):
            logger.error(f"API returned error: {result.get('error')}")
            return generate_review_fallback(cv_data, review_type)
        
        logger.info(f"CV analysis completed successfully. Score: {result.get('improvement_score')}")
        return result
    
    except Exception as e:
        logger.error(f"Error calling CV analysis API: {str(e)}")
        return generate_review_fallback(cv_data, review_type)

def generate_review_fallback(cv_data, review_type):
    """
    Generate a fallback review when API is unavailable
    
    Args:
        cv_data (dict): CV data with text and structure
        review_type (str): Type of review (basic or advanced)
        
    Returns:
        dict: Generated review results
    """
    if review_type == 'basic':
        return generate_basic_review_fallback(cv_data)
    else:
        return generate_advanced_review_fallback(cv_data)

def generate_basic_review_fallback(cv_data):
    """Generate a basic review fallback"""
    text = cv_data.get('full_text', '')
    sections = cv_data.get('sections', {}).get('sections', {})
    
    # Initialize insights list
    insights = []
    
    # Check for common CV sections
    if not sections.get('summary'):
        insights.append("Consider adding a professional summary at the top of your CV to highlight your key qualifications.")
    
    if not sections.get('experience'):
        insights.append("Your work experience section is missing or not clearly defined. This is a critical section for most CVs.")
    
    if not sections.get('education'):
        insights.append("Include your educational background with relevant details about degrees, institutions, and graduation dates.")
    
    if not sections.get('skills'):
        insights.append("A skills section would help highlight your key competencies relevant to your target roles.")
    
    # Check for contact info
    contact_info = cv_data.get('sections', {}).get('contact_info', {})
    if not contact_info.get('email') and not contact_info.get('phone'):
        insights.append("Ensure your contact information is clearly visible at the top of your CV.")
    
    # Check for quantifiable achievements
    if text.find('%') == -1 and text.find('increased') == -1 and text.find('decreased') == -1:
        insights.append("Add quantifiable achievements with metrics (%, numbers, etc.) to make your accomplishments more impactful.")
    
    # Check overall length
    words = text.split()
    if len(words) < 300:
        insights.append("Your CV appears to be quite short. Consider adding more details about your experience and skills.")
    elif len(words) > 1000:
        insights.append("Your CV may be too long. Consider condensing it to 1-2 pages for better readability.")
    
    # Add general insights if we don't have enough
    if len(insights) < 5:
        general_insights = [
            "Use action verbs at the beginning of bullet points to create a stronger impression.",
            "Ensure consistent formatting throughout your CV.",
            "Tailor your CV for each job application to highlight the most relevant experience.",
            "Consider the order of your sections for maximum impact.",
            "Proofread carefully for grammar and spelling errors."
        ]
        
        insights.extend(general_insights[:5 - len(insights)])
    
    return {
        'success': True,
        'timestamp': datetime.now().isoformat(),
        'insights': insights,
        'api_provider': 'Local Analysis (Fallback)'
    }

def generate_advanced_review_fallback(cv_data):
    """Generate an advanced review fallback"""
    text = cv_data.get('full_text', '')
    sections = cv_data.get('sections', {}).get('sections', {})
    
    # Calculate a score based on CV content
    score = 60  # Base score
    
    # Adjust score based on content
    if sections.get('summary'):
        score += 5
    if sections.get('experience'):
        score += 5
    if sections.get('education'):
        score += 5
    if sections.get('skills'):
        score += 5
    
    # Check for contact info
    contact_info = cv_data.get('sections', {}).get('contact_info', {})
    if contact_info.get('email') or contact_info.get('phone'):
        score += 5
    
    # Check for quantifiable achievements
    if text.find('%') != -1 or text.find('increased') != -1 or text.find('decreased') != -1:
        score += 5
    
    # Cap score at 90
    score = min(score, 90)
    
    # Generate detailed insights
    insights = [
        f"STRUCTURE: Your CV {'includes most key sections' if score > 70 else 'is missing some important sections'}. {'' if score > 75 else 'Consider reorganizing to include all essential information.'}",
        
        "CONTENT: Your achievements need more specific metrics. Add numbers, percentages, and outcomes to make your impact clear.",
        
        f"SKILLS: {'Your technical skills section is well organized.' if sections.get('skills') else 'Add a dedicated skills section to highlight your technical and soft skills.'} Group similar skills and indicate proficiency levels.",
        
        "LANGUAGE: Use more action verbs at the beginning of bullet points to create a stronger impression.",
        
        "FORMATTING: Maintain consistent formatting with a clear hierarchy of no more than 3 font sizes.",
        
        "RELEVANCE: Focus on your most recent and relevant experience. Older positions can be summarized briefly.",
        
        f"EDUCATION: {'Your education section is well-structured.' if sections.get('education') else 'Add an education section with degrees, institutions, and graduation dates.'}",
        
        "KEYWORDS: Include industry-specific keywords to pass through applicant tracking systems.",
        
        f"CONTACT INFO: {'Your contact information is present,' if contact_info.get('email') or contact_info.get('phone') else 'Add your contact information'} at the top of your CV for easy access.",
        
        "LENGTH: Keep your CV to 1-2 pages for optimal readability."
    ]
    
    return {
        'success': True,
        'timestamp': datetime.now().isoformat(),
        'improvement_score': score,
        'insights': insights,
        'api_provider': 'Local Analysis (Fallback)'
    }

def generate_pdf_report(review_result, cv_path):
    """
    Generate a PDF report from review results
    
    Args:
        review_result (dict): Review results
        cv_path (str): Path to the original CV
        
    Returns:
        str: Path to the generated PDF report
    """
    # Create a temporary file for the report
    timestamp = int(time.time())
    report_filename = f"cv_review_report_{timestamp}.pdf"
    report_path = os.path.join(Config.UPLOAD_FOLDER, report_filename)
    os.makedirs(Config.UPLOAD_FOLDER, exist_ok=True)
    
    # Create the PDF document
    doc = SimpleDocTemplate(
        report_path,
        pagesize=letter,
        rightMargin=72,
        leftMargin=72,
        topMargin=72,
        bottomMargin=72
    )
    
    # Get styles
    styles = getSampleStyleSheet()
    
    # Create custom styles
    title_style = ParagraphStyle(
        'Title',
        parent=styles['Title'],
        fontSize=24,
        alignment=1,  # Center alignment
        spaceAfter=24
    )
    
    heading_style = ParagraphStyle(
        'Heading1',
        parent=styles['Heading1'],
        fontSize=18,
        spaceBefore=12,
        spaceAfter=6
    )
    
    subheading_style = ParagraphStyle(
        'Heading2',
        parent=styles['Heading2'],
        fontSize=14,
        spaceBefore=12,
        spaceAfter=6
    )
    
    body_style = ParagraphStyle(
        'Body',
        parent=styles['Normal'],
        fontSize=12,
        spaceAfter=12
    )
    
    # Build the PDF content
    content = []
    
    # Title
    content.append(Paragraph("CV Review Report", title_style))
    content.append(Spacer(1, 12))
    
    # Date
    date_text = f"Generated on: {datetime.now().strftime('%d %B %Y')}"
    content.append(Paragraph(date_text, styles['Italic']))
    content.append(Spacer(1, 24))
    
    # Score section
    content.append(Paragraph("CV Improvement Score", heading_style))
    score = review_result.get('improvement_score', 0)
    content.append(Paragraph(f"{score}/100", subheading_style))
    content.append(Spacer(1, 12))
    
    # Score interpretation
    if score < 50:
        score_text = "Your CV needs significant improvement to be competitive in the job market."
    elif score < 70:
        score_text = "Your CV is adequate but could benefit from several improvements."
    else:
        score_text = "Your CV is strong, with only minor improvements needed."
    
    content.append(Paragraph(score_text, body_style))
    content.append(Spacer(1, 24))
    
    # Insights section
    content.append(Paragraph("Detailed Insights", heading_style))
    content.append(Spacer(1, 12))
    
    insights = review_result.get('insights', [])
    
    # Check if insights have category prefixes (like "STRUCTURE: ")
    if any(':' in insight for insight in insights):
        # Group insights by category
        categories = {}
        
        for insight in insights:
            if ':' in insight:
                category, detail = insight.split(':', 1)
                category = category.strip()
                detail = detail.strip()
                
                if category not in categories:
                    categories[category] = []
                
                categories[category].append(detail)
            else:
                if 'GENERAL' not in categories:
                    categories['GENERAL'] = []
                
                categories['GENERAL'].append(insight)
        
        # Add each category and its insights
        for category, category_insights in categories.items():
            content.append(Paragraph(category, subheading_style))
            
            bullets = []
            for detail in category_insights:
                bullets.append(ListItem(Paragraph(detail, body_style)))
            
            content.append(ListFlowable(
                bullets,
                bulletType='bullet',
                leftIndent=20
            ))
            
            content.append(Spacer(1, 12))
    else:
        # Simple list of insights
        bullets = []
        for insight in insights:
            bullets.append(ListItem(Paragraph(insight, body_style)))
        
        content.append(ListFlowable(
            bullets,
            bulletType='bullet',
            leftIndent=20
        ))
    
    content.append(Spacer(1, 24))
    
    # Next steps section
    content.append(Paragraph("Next Steps", heading_style))
    content.append(Spacer(1, 6))
    
    next_steps = [
        "Review the insights and prioritize areas for improvement",
        "Update your CV based on the recommendations",
        "Tailor your CV for each job application",
        "Ask a colleague or mentor to review your updated CV",
        "Consider professional design services if formatting is an issue"
    ]
    
    bullets = []
    for step in next_steps:
        bullets.append(ListItem(Paragraph(step, body_style)))
    
    content.append(ListFlowable(
        bullets,
        bulletType='number',
        leftIndent=20
    ))
    
    # Build the PDF
    doc.build(content)
    
    return report_path

def get_phone_number_from_storage_path(storage_path):
    """
    Extract phone number from storage path
    
    Args:
        storage_path (str): Storage path
        
    Returns:
        str: Phone number
    """
    # Path format: cv-uploads/{phoneNumber}/{filename}
    parts = storage_path.split('/')
    if len(parts) >= 2:
        return parts[1]
    return None